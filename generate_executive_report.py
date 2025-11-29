"""
Script para generar un reporte ejecutivo en Word
a partir de los resultados del backtest.

Usa:
    - outputs/backtest/summary_best_runs.(csv|xlsx)
    - outputs/backtest/report_ARIMA.csv, report_LSTM.csv, etc. (opcional)
    - outputs/predictions/production_signals.(csv|xlsx) (opcional)

Salida:
    outputs/reportes/reporte_ejecutivo_modelos.docx
"""

import os
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ==========================
# CONFIGURACIÓN BÁSICA
# ==========================

PROJECT_NAME = "Prototipo de sistema automatizado de inversión"
ASSET_UNIVERSE = "EUR/USD y SPY"   # cámbialo si solo usas EUR/USD, etc.

BASE_DIR = Path(__file__).parent
BACKTEST_DIR = BASE_DIR / "outputs" / "backtest"
PREDICTIONS_DIR = BASE_DIR / "outputs" / "predictions"
REPORTS_DIR = BASE_DIR / "outputs" / "reportes"

SUMMARY_PATHS = [
    BACKTEST_DIR / "summary_best_runs.csv",
    BACKTEST_DIR / "summary_best_runs.xlsx"
]

MODEL_REPORTS = {
    "ARIMA": BACKTEST_DIR / "report_ARIMA.csv",
    "LSTM": BACKTEST_DIR / "report_LSTM.csv",
    "PROPHET": BACKTEST_DIR / "report_PROPHET.csv",
    "RandomWalk": BACKTEST_DIR / "report_RandomWalk.csv",
}

PRODUCTION_SIGNALS_PATHS = [
    PREDICTIONS_DIR / "production_signals.csv",
    PREDICTIONS_DIR / "production_signals.xlsx",
]

OUTPUT_DOCX = REPORTS_DIR / "reporte_ejecutivo_modelos.docx"


# ==========================
# FUNCIONES AUXILIARES
# ==========================

def load_table(path_candidates):
    """
    Intenta cargar un archivo CSV o Excel de una lista de rutas candidatas.
    Devuelve un DataFrame o None si nada se puede cargar.
    """
    for path in path_candidates:
        if path.exists():
            suffix = path.suffix.lower()
            if suffix in [".csv"]:
                return pd.read_csv(path)
            elif suffix in [".xlsx", ".xls"]:
                return pd.read_excel(path)
    return None


def format_float(value, decimals=4):
    """Formatea un valor numérico con n decimales; si falla, devuelve '-'."""
    try:
        return f"{float(value):.{decimals}f}"
    except Exception:
        return "-"


def add_title(document: Document, text: str):
    """Añade título principal con estilo grande."""
    title = document.add_heading(text, level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(document: Document, text: str, level: int = 2):
    """Añade encabezado de sección."""
    document.add_heading(text, level=level)


def add_paragraph(document: Document, text: str, bold=False, italic=False):
    """Añade un párrafo con formato sencillo."""
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(6)
    return p


def create_summary_table(document: Document, df_summary: pd.DataFrame):
    """
    Crea una tabla en Word con resumen por modelo (RMSE, MAE, Hit Rate).
    Toma los datos de summary_best_runs.
    """
    required_cols = ["model"]
    metrics_cols = [col for col in ["rmse", "mae", "hit_rate"] if col in df_summary.columns]
    cols = required_cols + metrics_cols

    df = df_summary[cols].copy()

    # Crear tabla
    table = document.add_table(rows=len(df) + 1, cols=len(cols))
    table.style = "Table Grid"

    # Encabezados
    header_cells = table.rows[0].cells
    header_cells[0].text = "Modelo"
    for i, metric in enumerate(metrics_cols, start=1):
        if metric == "hit_rate":
            header_cells[i].text = "Hit Rate (%)"
        else:
            header_cells[i].text = metric.upper()

    # Filas
    for row_idx, (_, row) in enumerate(df.iterrows(), start=1):
        row_cells = table.rows[row_idx].cells
        row_cells[0].text = str(row["model"])
        for i, metric in enumerate(metrics_cols, start=1):
            row_cells[i].text = format_float(row[metric])

    return table


# ==========================
# GENERACIÓN DEL REPORTE
# ==========================

def generate_executive_report():
    # Crear directorio de salida
    os.makedirs(REPORTS_DIR, exist_ok=True)

    # Cargar resumen de mejores ejecuciones
    df_summary = load_table(SUMMARY_PATHS)
    if df_summary is None:
        raise FileNotFoundError(
            "No se encontró summary_best_runs en outputs/backtest "
            "(ni CSV ni XLSX). Ejecuta el backtest primero."
        )

    # Cargar señales de producción (opcional)
    df_prod = load_table(PRODUCTION_SIGNALS_PATHS)

    # Crear documento
    doc = Document()

    # 0. Título / portada
    add_title(doc, "Reporte ejecutivo de evaluación de modelos de trading")

    add_paragraph(
        doc,
        f"Proyecto: {PROJECT_NAME}\nActivos analizados: {ASSET_UNIVERSE}",
        italic=True
    )

    # 1. Objetivo del reporte
    add_heading(doc, "1. Objetivo del reporte", level=2)
    add_paragraph(
        doc,
        "El objetivo de este reporte es presentar, en lenguaje ejecutivo, "
        "los resultados de la evaluación de modelos de predicción utilizados "
        "en el prototipo de sistema automatizado de inversión. "
        "Se resumen las métricas de desempeño de los modelos ARIMA, Prophet, "
        "LSTM y del benchmark Random Walk/Momentum, calculadas mediante un "
        "esquema de backtesting tipo walk-forward."
    )

    # 2. Descripción general del experimento
    add_heading(doc, "2. Descripción general del experimento", level=2)
    add_paragraph(
        doc,
        "El experimento se basa en un pipeline que integra: (i) carga de datos "
        "históricos desde MetaTrader 5, (ii) limpieza y validación de datos OHLC, "
        "(iii) generación de retornos e indicadores técnicos, y (iv) evaluación "
        "de modelos mediante backtesting walk-forward."
    )

    # 3. Modelos evaluados
    add_heading(doc, "3. Modelos evaluados", level=2)
    add_paragraph(doc, "• Random Walk / Momentum (benchmark).")
    add_paragraph(doc, "• ARIMA con regresión Ridge sobre los residuos.")
    add_paragraph(doc, "• Prophet con regresores exógenos opcionales.")
    add_paragraph(doc, "• LSTM (red recurrente para series temporales).")

    # 4. Métricas de evaluación
    add_heading(doc, "4. Métricas de evaluación", level=2)
    add_paragraph(doc, "• RMSE (Root Mean Squared Error).")
    add_paragraph(doc, "• MAE (Mean Absolute Error).")
    add_paragraph(doc, "• Hit Rate (% de aciertos en la dirección del retorno).")

    # 5. Resultados principales
    add_heading(doc, "5. Resultados principales", level=2)
    add_paragraph(doc, "5.1 Resumen de desempeño por modelo", bold=True)

    # Tabla de resumen
    create_summary_table(doc, df_summary)

    # Identificar el mejor modelo global por RMSE
    if "rmse" in df_summary.columns:
        best_row = df_summary.loc[df_summary["rmse"].idxmin()]
        best_model = best_row["model"]
        txt_best = (
            f"El modelo con mejor desempeño según RMSE es {best_model}, "
            f"con RMSE={format_float(best_row.get('rmse'))}, "
            f"MAE={format_float(best_row.get('mae'))} y "
            f"Hit Rate={format_float(best_row.get('hit_rate'))} %."
        )
        add_paragraph(doc, txt_best)

    # Espacios para gráficos
    add_paragraph(doc, "")
    add_paragraph(doc, "5.2 Comparación ARIMA vs Random Walk (benchmark)", bold=True)
    add_paragraph(doc, "[Espacio reservado para insertar el gráfico ARIMA vs Random Walk].")

    add_paragraph(doc, "")
    add_paragraph(doc, "5.3 Desempeño de LSTM y Prophet", bold=True)
    add_paragraph(doc, "[Espacio reservado para insertar la comparación gráfica LSTM / Prophet / ARIMA].")

    # 6. Producción (opcional)
    add_heading(doc, "6. Señales de producción (opcional)", level=2)
    if df_prod is not None:
        add_paragraph(
            doc,
            f"Se cargaron {len(df_prod)} registros de señales en producción "
            "desde el archivo production_signals. A partir de este conjunto "
            "se pueden calcular métricas de trading (número de operaciones, "
            "proporción de BUY/SELL/HOLD, rendimiento acumulado, etc.)."
        )
    else:
        add_paragraph(
            doc,
            "No se encontraron señales de producción. Esta sección puede "
            "completarse una vez se ejecuten escenarios en modo producción."
        )

    # 7. Implicaciones y conclusiones
    add_heading(doc, "7. Implicaciones para la toma de decisiones", level=2)
    add_paragraph(
        doc,
        "Los resultados muestran que los modelos evaluados pueden ofrecer "
        "una mejora marginal sobre el benchmark aleatorio en términos de "
        "error de predicción y acierto direccional. Sin embargo, dichas "
        "ganancias deben ponderarse frente a los costos de transacción y "
        "la complejidad operativa del modelo."
    )

    # 8. Limitaciones y trabajo futuro
    add_heading(doc, "8. Limitaciones y trabajo futuro", level=2)
    add_paragraph(
        doc,
        "El periodo de evaluación es limitado y puede no capturar todos los "
        "regímenes de mercado. Como trabajo futuro se propone: (i) extender "
        "la evaluación a otros activos y horizontes temporales, (ii) incluir "
        "métricas de retorno ajustado por riesgo, y (iii) analizar la "
        "sensibilidad del sistema a cambios en los hiperparámetros y reglas de decisión."
    )

    # Guardar documento
    doc.save(OUTPUT_DOCX)
    print(f"✅ Reporte ejecutivo generado en: {OUTPUT_DOCX}")


if __name__ == "__main__":
    generate_executive_report()
